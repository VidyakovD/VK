import io
import re
from pathlib import Path

import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.services.llm.base import build_httpx_client


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".md"}


def parse_document_bytes(data: bytes, filename: str) -> str:
    """
    Extract plain text from a document. Dispatches by file extension.
    """
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(data)
    if ext == ".docx":
        return _parse_docx(data)
    if ext in (".txt", ".csv", ".md"):
        return data.decode("utf-8", errors="replace")
    raise ValueError(f"Unsupported file extension: {ext}")


def _parse_pdf(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    parts = []
    for page in reader.pages:
        text = page.extract_text() or ""
        parts.append(text)
    return _normalize_whitespace("\n\n".join(parts))


def _parse_docx(data: bytes) -> str:
    doc = DocxDocument(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Tables (best effort)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                parts.append(row_text)
    return _normalize_whitespace("\n".join(parts))


async def parse_url(url: str, timeout: float = 30.0) -> str:
    """Fetch URL and extract readable text (basic, no JS execution)."""
    async with build_httpx_client(timeout=timeout) as client:
        resp = await client.get(url, follow_redirects=True)
        resp.raise_for_status()
        ctype = resp.headers.get("content-type", "")
        if "text/html" in ctype or "<html" in resp.text[:500].lower():
            return _extract_html_text(resp.text)
        return _normalize_whitespace(resp.text)


def _extract_html_text(html: str) -> str:
    soup = BeautifulSoup(html, "lxml")
    for tag in soup(["script", "style", "noscript", "nav", "header", "footer", "aside", "form"]):
        tag.decompose()
    main = soup.find("main") or soup.find("article") or soup.body or soup
    text = main.get_text(separator="\n")
    return _normalize_whitespace(text)


_MULTI_WHITESPACE = re.compile(r"[ \t]+")
_MULTI_NEWLINE = re.compile(r"\n{3,}")


def _normalize_whitespace(text: str) -> str:
    text = _MULTI_WHITESPACE.sub(" ", text)
    text = _MULTI_NEWLINE.sub("\n\n", text)
    return text.strip()
